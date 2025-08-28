from otree.api import *  # Core oTree framework
from .models import C, get_vacancy_info, get_applicants_data_for_vacancy, \
    load_metadata_criteria, should_show_vacancy_session, get_applicant_ids, \
    assign_static_role  # imports from models.py
import random  # for StroopTest Items
from docx import Document  # Word -> HTML converting
import os  # file paths
import json


class Consent(Page):

    def is_displayed(self):
        return self.player.round_number == 1  # only shown in the very first round


# MAIN TASK PAGES

class Recruiter(Page):
    """
    Reviews applicant information including CVs, job references, and cover letters.
    Now used in the 5-round structure for all vacancies

    Data Processing Flow:
    1. Get current vacancy configuration
    2. Load basic applicant data (names, document paths)
    3. Load and convert Word documents to HTML for each applicant
    4. Calculate session numbering and progress indicators
    5. Assemble all data for template rendering

    Returns:
    dict: Template variables containing:
        - applicants: List of applicant objects with HTML content
        - session_number: Current session
        - vacancy_number: Which vacancy is active
        - remaining_time: Session timeout in seconds
        - static_path: Path to static files (PDFs, images)
        - total_sessions: Total number of working sessions (3)
    """

    def is_displayed(self):
        """
        Shows recruiter interface only during vacancy rounds with role assignment.
        """
        if not should_show_vacancy_session(self.player.round_number):
            return False

        # AUTOMATIC ROLE ASSIGNMENT
        assign_static_role(self.player)
        return self.player.is_recruiter()

    def get_timeout_seconds(self):
        """
        Returns timeout based on vacancy
        """
        if self.player.round_number == C.VACANCY_1_ROUND:
            return None  # Unlimited time for Vacancy 1

        vacancy_info = get_vacancy_info(self.player.round_number, self.player)
        return vacancy_info['duration_seconds'] if vacancy_info else 600  # 10 min fallback

    timeout_seconds = property(get_timeout_seconds)  # Convert method to property for oTree

    def vars_for_template(self):
        """
        Loads and processes all data needed for the recruiter interface.
        """
        vacancy_info = get_vacancy_info(self.player.round_number, self.player)
        applicants_data = get_applicants_data_for_vacancy(vacancy_info)

        applicants_with_content = []
        doc_suffix = vacancy_info['doc_suffix'] if vacancy_info else '1'

        for applicant in applicants_data:
            # Create copy to avoid modifying original data
            applicant_data = applicant.copy()
            applicant_data['description'] = self.get_word_content(applicant['id'], doc_suffix)
            applicants_with_content.append(applicant_data)

        vacancy_number = vacancy_info['vacancy'] if vacancy_info else 1

        return {
            'applicants': applicants_with_content,
            'session_number': vacancy_number,
            'vacancy_number': vacancy_number,
            'remaining_time': vacancy_info['duration_seconds'] if vacancy_info else 600,
            'static_path': C.STATIC_APPLICANTS_PATH,
            'total_sessions': 3  # 2 working sessions total
        }

    def get_word_content(self, applicant_id, doc_suffix='1'):
        """
        Loads recruiter mask Word document.

        Args:
            applicant_id (str): Applicant identifier ('a', 'b', or 'c')
            doc_suffix (str): Vacancy-specific suffix

        Returns:
            str: HTML content for web display, or error message if file not found
        """
        try:
            # Build absolute path to Word document
            current_dir = os.path.dirname(os.path.abspath(__file__))

            doc_path = os.path.join(
                current_dir,
                '..',
                '_static',
                'applicants',
                f'recruiter_maske_{applicant_id}{doc_suffix}.docx'
            )

            # Normalize path for cross-platform compatibility
            doc_path = os.path.normpath(doc_path)

            # Load Word document and convert to HTML
            document = Document(doc_path)
            html_content = self.convert_docx_to_html(document)
            return html_content

        except Exception as e:
            return f"<p><em>Error loading document: {str(e)}</em></p>"

    def convert_docx_to_html(self, document):
        """
        Processes Word document content and converts it to styled HTML.

        Args:
        document: python-docx Document object

        Returns:
        str: HTML content with styling ready for web display
        """
        html_parts = []

        # Process paragraphs and headings
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                # Convert Word headings to HTML heading tags
                if paragraph.style.name.startswith('Heading'):
                    level = 3
                    if 'Heading 1' in paragraph.style.name:
                        level = 2
                    elif 'Heading 2' in paragraph.style.name:
                        level = 3
                    html_parts.append(f'<h{level}>{text}</h{level}>')
                else:
                    # Regular paragraphs with bold text detection
                    if any(run.bold for run in paragraph.runs):
                        html_parts.append(f'<p><strong>{text}</strong></p>')
                    else:
                        html_parts.append(f'<p>{text}</p>')

        # Process tables with styling and header detection
        for table in document.tables:
            html_parts.append('<table style="width:100%; border-collapse: collapse; margin: 10px 0;">')

            for row_idx, row in enumerate(table.rows):
                html_parts.append('<tr>')
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    # Style header cells differently
                    if row_idx == 0 or any(run.bold for para in cell.paragraphs for run in para.runs):
                        html_parts.append(
                            f'<td style="border: 1px solid #ddd; padding: 5px; font-weight: bold; background-color: #f5f5f5;">{cell_text}</td>')
                    else:
                        html_parts.append(f'<td style="border: 1px solid #ddd; padding: 5px;">{cell_text}</td>')
                html_parts.append('</tr>')

            html_parts.append('</table>')

        if not html_parts:
            return "<p><em>The Word document is empty or could not be read.</em></p>"

        # Combine all HTML parts into final output
        return ''.join(html_parts)


class HRCoordinator(Page):
    """
    Prepares all data needed for HR Coordinator interface.
    Now used in the 5-round structure with predefined criteria auto-loading.

    Data Processing Flow:
    1. Get current vacancy configuration
    2. Load basic applicant data (names, document paths)
    3. Load evaluation metadata for interactive criteria system
    4. Calculate session numbering and progress indicators
    5. Assemble evaluation interface data for template rendering

    Returns:
        dict: Template variables containing:
            - applicants: Full applicant data for evaluation table
            - criteria_data: All criteria for modal selection system
            - predefined_criteria: Criteria that should be auto-loaded
            - categories/criteria_by_category: Criteria organization for modal
            - relevance_factors: Scoring multipliers (low:1, normal:2, high:3)
            - applicant_colors: Colors for live pie chart display
            - job_desc_file: PDF file for job description access
            - min_score/max_score: Score range for evaluation dropdowns (0-8)
            - applicant_ids: List for score validation (['a', 'b', 'c'])

    Form Fields:
        - criteria_added_this_session: Count of criteria added by player
        - validation_data_json: Hidden field containing criteria data for validation
    """
    form_model = 'player'
    form_fields = ['criteria_added_this_session', 'validation_data_json']

    def is_displayed(self):
        """
        Display logic and automatic role assignment.
        Uses static role assignment for 4-round structure.
        """
        if not should_show_vacancy_session(self.player.round_number):
            return False

        # AUTOMATIC ROLE ASSIGNMENT
        assign_static_role(self.player)
        return self.player.is_hr_coordinator()

    def get_timeout_seconds(self):
        """
        Returns timeout based on vacancy: None for V1 (unlimited), 600s for V2.
        """
        if self.player.round_number == C.VACANCY_1_ROUND:
            return None  # Unlimited time for Vacancy 1

        vacancy_info = get_vacancy_info(self.player.round_number, self.player)
        return vacancy_info['duration_seconds'] if vacancy_info else 600  # 10 min fallback

    timeout_seconds = property(get_timeout_seconds)

    def before_next_page(self):
        """
        Validates player's criteria data against correct answers before proceeding.
        Takes JSON data from frontend, compares it with metadata, and counts how many criteria were evaluated correctly vs incorrectly.
        """

        try:
            # Get criteria data from hidden form field (sent by JavaScript)
            validation_data_str = self.player.validation_data_json or '{}'

            if not validation_data_str or validation_data_str == '{}':
                self.player.criteria_correct_this_session = 0
                self.player.criteria_incorrect_this_session = 0
                return

            # Parse JSON data from frontend
            criteria_data = json.loads(validation_data_str)
            correct_count = 0
            incorrect_count = 0

            # Load correct answers from metadata files
            metadata = load_metadata_criteria(self.player.round_number, self.player)
            metadata_criteria = metadata['criteria']

            # Check each criterion the player evaluated
            for criterion_name, data in criteria_data.items():
                # Find the criterion in metadata
                criterion_metadata = None
                for criterion in metadata_criteria:
                    if criterion['name'].strip().lower() == criterion_name.strip().lower():
                        criterion_metadata = criterion
                        break

                if criterion_metadata:
                    # Check scores and relevance
                    scores_correct = True

                    # Check scores for each applicant
                    for applicant_id in get_applicant_ids():
                        entered_score = data['scores'].get(applicant_id, 0)
                        correct_score = criterion_metadata['scores'].get(f'applicant_{applicant_id}', 0)

                        if int(entered_score or 0) != int(correct_score):
                            scores_correct = False
                            break

                    # Validate relevance level
                    entered_relevance = data.get('relevance', 'normal')
                    correct_relevance = criterion_metadata.get('relevance', 'normal')
                    relevance_correct = entered_relevance == correct_relevance

                    # Count correct or incorrect
                    if scores_correct and relevance_correct:
                        correct_count += 1
                    else:
                        incorrect_count += 1
                else:
                    # Criterion not found in metadata = incorrect
                    incorrect_count += 1

            # Save results to player fields for later analysis
            self.player.criteria_correct_this_session = correct_count
            self.player.criteria_incorrect_this_session = incorrect_count

        except Exception as e:
            self.player.criteria_correct_this_session = 0
            self.player.criteria_incorrect_this_session = 0

    def vars_for_template(self):
        """
        Prepares all data needed for HR Coordinator interface.
        """
        vacancy_info = get_vacancy_info(self.player.round_number, self.player)
        applicants_data = get_applicants_data_for_vacancy(vacancy_info)

        # Load evaluation criteria and categories from Excel metadata
        metadata = load_metadata_criteria(self.player.round_number, self.player)

        vacancy_number = vacancy_info['vacancy'] if vacancy_info else 1

        return {
            'applicants': applicants_data,
            'min_score': C.MIN_SCORE,
            'max_score': C.MAX_SCORE,
            'session_number': vacancy_number,
            'vacancy_number': vacancy_number,
            'remaining_time': vacancy_info['duration_seconds'] if vacancy_info else 600,
            'criteria_data': metadata['criteria'],
            'predefined_criteria': metadata['predefined_criteria'],
            'categories': metadata['categories'],
            'criteria_by_category': metadata['criteria_by_category'],
            'relevance_factors': C.RELEVANCE_FACTORS,
            'job_desc_file': vacancy_info['job_desc_file'] if vacancy_info else 'job_description_1.pdf',
            'static_path': C.STATIC_APPLICANTS_PATH,
            'applicant_colors': C.APPLICANT_COLORS,
            'applicant_ids': get_applicant_ids(),
            'total_sessions': 3
        }


class BusinessPartner(Page):
    """
    Prepares data for Business Partner requirements catalog interface.

    Data Processing Flow:
    1. Get current vacancy configuration
    2. Load basic applicant data (not used directly in interface)
    3. Load criteria metadata for requirements catalog
    4. Calculate session numbering and progress indicators
    5. Generate vacancy-specific file paths for emails and sticky notes
    6. Assemble catalog data for template rendering

    Returns:
        dict: Template variables containing:
            - criteria_data: All evaluation criteria for catalog browsing
            - categories: Criteria categories for navigation
            - criteria_by_category: Grouped criteria for category selection
            - min_score/max_score: Score range for criteria viewing (0-8)
            - email_file: Vacancy-specific email PDF filename
            - sticky_notes_file: Vacancy-specific sticky notes image filename
    """

    def is_displayed(self):
        if should_show_vacancy_session(self.player.round_number):
            assign_static_role(self.player)
            return self.player.is_business_partner()
        return False

    def get_timeout_seconds(self):
        """
        Returns timeout based on vacancy: None for V1 (unlimited), 600s for V2 and V3.
        """
        if self.player.round_number == C.VACANCY_1_ROUND:
            return None  # Unlimited time for Vacancy 1

        vacancy_info = get_vacancy_info(self.player.round_number, self.player)
        return vacancy_info['duration_seconds'] if vacancy_info else 600  # 10 min fallback

    timeout_seconds = property(get_timeout_seconds)

    def vars_for_template(self):
        """
        Prepares data for Business Partner requirements catalog interface.
        Now includes vacancy-specific email and sticky notes files.
        """
        vacancy_info = get_vacancy_info(self.player.round_number, self.player)
        applicants_data = get_applicants_data_for_vacancy(vacancy_info)

        metadata = load_metadata_criteria(self.player.round_number, self.player)

        vacancy_number = vacancy_info['vacancy'] if vacancy_info else 1

        return {
            'applicants': applicants_data,
            'session_number': vacancy_number,
            'vacancy_number': vacancy_number,
            'remaining_time': vacancy_info['duration_seconds'] if vacancy_info else 600,
            'criteria_data': metadata['criteria'],
            'categories': metadata['categories'],
            'criteria_by_category': metadata['criteria_by_category'],
            'static_path': C.STATIC_APPLICANTS_PATH,
            'min_score': C.MIN_SCORE,
            'max_score': C.MAX_SCORE,
            'total_sessions': 3,
            'email_file': f'Email_{vacancy_number}.pdf',
            'sticky_notes_file': f'StickyNotes_{vacancy_number}.jpg'
        }


class WaitForVacancy(WaitPage):
    """
    Waits for all players before a vacancy session starts.
    Ensures that all roles (Recruiter, HR Coordinator, Business Partner)
    begin the work session simultaneously.
    """

    def is_displayed(self):
        """
        Only displayed before main work sessions (Vacancy 1, 2, 3).
        """
        return should_show_vacancy_session(self.player.round_number)

    def after_all_players_arrive(self):
        """
        Executed when all players have arrived.
        Add setup code here if needed.
        """
        pass

    title_text = "Synchronization"
    body_text = "Waiting for other players to join the session..."


class SelfAssessment(Page):
    """
    Post-session questionnaire for measuring mental fatigue and subjective workload (on a scale of 1 - 10).
    Now shown in all 3 rounds: Baseline (Round 1), After V1 (Round 2), After V2 (Round 3).

    Form Fields:
        - fatigue_level
        - mental_effort
        - concentration_difficulty
        - motivation_level
    """
    form_model = 'player'
    form_fields = ['fatigue_level', 'mental_effort', 'concentration_difficulty', 'motivation_level']

    def is_displayed(self):
        """
        Shown in all 3 measurement rounds: Baseline (1), After V1 (2), After V2 (3).
        """
        return self.player.round_number in [C.CONSENT_ROUND, C.VACANCY_1_ROUND, C.VACANCY_2_ROUND, C.VACANCY_3_ROUND]

    def vars_for_template(self):
        """
        Prepares session information for self-assessment form.
        """
        vacancy_info = get_vacancy_info(self.player.round_number, self.player)
        vacancy_number = vacancy_info['vacancy'] if vacancy_info else 0  # 0 for baseline

        # Determine session description
        if self.player.round_number == C.CONSENT_ROUND:
            session_name = "Baseline"
            session_number = 0
        else:
            session_name = f"Vacancy {vacancy_number}"
            session_number = vacancy_number

        return {
            'session_number': session_number,
            'session_name': session_name,
            'vacancy_number': vacancy_number,
            'role_played': self.player.selected_role,
            'total_sessions': 3
        }


class CognitiveTestInstructions(Page):
    """
    Instructions page for cognitive test before actual test execution.
    Now shown in all 3 measurement rounds.
    """
    timeout_seconds = 20

    def is_displayed(self):
        """
        Only shown before the first cognitive test.
        """
        return self.player.round_number == C.CONSENT_ROUND

    def vars_for_template(self):
        """
        Prepares session information for test instructions.
        """
        vacancy_info = get_vacancy_info(self.player.round_number, self.player)
        vacancy_number = vacancy_info['vacancy'] if vacancy_info else 0  # 0 for baseline

        # Determine session description
        if self.player.round_number == C.CONSENT_ROUND:
            session_name = "Baseline"
            session_number = 0
        else:
            session_name = f"Vacancy {vacancy_number}"
            session_number = vacancy_number

        return {
            'session_number': session_number,
            'session_name': session_name,
            'vacancy_number': vacancy_number,
            'total_sessions': 3
        }


class CognitiveTest(Page):
    """
    Stroop test for measuring cognitive load and performance changes over sessions.
    Now executed in all 3 measurement rounds for baseline comparison.

    Data Processing Flow:
    1. Get current vacancy configuration
    2. Generate random Stroop test items (word-color mismatches)
    3. Create color mapping for answer validation
    4. Calculate session numbering and progress indicators
    5. Assemble test data for interactive interface

    Returns:
    dict: Template variables containing:
        - test_items: List of 20 random Stroop test items
        - test_duration: Test time limit (22 seconds)

    Form Fields:
        - cognitive_test_score: Number of correct answers
        - cognitive_test_reaction_time: Average reaction time in milliseconds
        - cognitive_test_errors: Number of incorrect answers
    """
    form_model = 'player'
    form_fields = ['cognitive_test_score', 'cognitive_test_reaction_time', 'cognitive_test_errors']
    timeout_seconds = None

    def is_displayed(self):
        """
        Shown in all 3 measurement rounds: Baseline (1), After V1 (2), After V2 (3).
        """
        return self.player.round_number in [C.CONSENT_ROUND, C.VACANCY_1_ROUND, C.VACANCY_2_ROUND, C.VACANCY_3_ROUND]

    def vars_for_template(self):
        """
        Generates random Stroop test items and prepares test interface.
        """
        vacancy_info = get_vacancy_info(self.player.round_number, self.player)
        vacancy_number = vacancy_info['vacancy'] if vacancy_info else 0  # 0 for baseline

        # Determine session description
        if self.player.round_number == C.CONSENT_ROUND:
            session_name = "Baseline"
            session_number = 0
        else:
            session_name = f"Vacancy {vacancy_number}"
            session_number = vacancy_number

        # Generate random Stroop test items
        test_items = []
        color_mapping = {
            '#ff0000': 'red',
            '#0000ff': 'blue',
            '#00ff00': 'green',
            '#ffff00': 'yellow'
        }

        for i in range(C.COGNITIVE_TEST_TOTAL_QUESTIONS):
            word = random.choice(C.STROOP_WORDS)
            color_hex = random.choice(C.STROOP_COLORS)
            color_name = color_mapping[color_hex]

            test_items.append({
                'word': word,
                'color': color_hex,
                'correct_answer': color_name
            })

        return {
            'test_items': test_items,
            'test_duration': C.COGNITIVE_TEST_DURATION,
            'session_number': session_number,
            'session_name': session_name,
            'vacancy_number': vacancy_number,
            'total_sessions': 3
        }


class CognitiveTestResults(Page):
    """
    Displays results of the completed Stroop test to participants.
    Now shown in all 3 measurement rounds.

    Returns:
        dict: Template variables containing:
            - score: Number of correct answers (0-20)
            - reaction_time: Average reaction time in milliseconds
            - errors: Number of incorrect answers
            - total_items: Total test items (20)
    """
    timeout_seconds = 20

    def is_displayed(self):
        """
        Shown in all 3 measurement rounds: Baseline (1), After V1 (2), After V2 (3).
        """
        return self.player.round_number in [C.CONSENT_ROUND, C.VACANCY_1_ROUND, C.VACANCY_2_ROUND, C.VACANCY_3_ROUND]

    def vars_for_template(self):
        """
        Loads cognitive test results and prepares results display.
        """
        vacancy_info = get_vacancy_info(self.player.round_number, self.player)
        vacancy_number = vacancy_info['vacancy'] if vacancy_info else 0  # 0 for baseline

        # Determine session description
        if self.player.round_number == C.CONSENT_ROUND:
            session_name = "Baseline"
            session_number = 0
        else:
            session_name = f"Vacancy {vacancy_number}"
            session_number = vacancy_number

        return {
            'session_number': session_number,
            'session_name': session_name,
            'vacancy_number': vacancy_number,
            'score': self.player.field_maybe_none('cognitive_test_score') or 0,
            'reaction_time': self.player.field_maybe_none('cognitive_test_reaction_time') or 0,
            'errors': self.player.field_maybe_none('cognitive_test_errors') or 0,
            'total_questions': C.COGNITIVE_TEST_TOTAL_QUESTIONS,
            'total_sessions': 3
        }


class FinalResults(Page):
    """
    Comprehensive results analysis comparing Baseline, Vacancy 1, and Vacancy 2.

    Analyzes 4 measurement points:
    - Baseline (Round 1): Before any work tasks
    - Vacancy 1 (Round 2): After unlimited time session
    - Vacancy 2 (Round 3): After time-limited session
    - Vacancy 3 (Round 4): After time-limited session

    Returns:
        dict: Template variables containing:
            - all_sessions: Complete session-by-session data (3 points)
            - baseline_* : Baseline measurements
            - v1_*_change: Changes from baseline to V1
            - v2_*_change: Changes from baseline to V2
            - trend data and statistical summaries
            - is_hr_coordinator: Boolean flag for role-specific content
    """

    def is_displayed(self):
        """
        Shown only in final round (Round 5).
        """
        return self.player.round_number == C.FINAL_RESULTS_ROUND

    # Final Results Page - Updated vars_for_template() method
    def vars_for_template(self):
        """
        Compiles results with separate baseline and task progression analysis
        """

        # Separate baseline from task sessions
        baseline_round = C.CONSENT_ROUND
        task_rounds = [C.VACANCY_1_ROUND, C.VACANCY_2_ROUND, C.VACANCY_3_ROUND]

        # Extract baseline data separately
        try:
            baseline_player = self.player.in_round(baseline_round)
            baseline_data = {
                'fatigue_level': baseline_player.fatigue_level or 0,
                'mental_energy': baseline_player.mental_effort or 0,  # "energy" in baseline
                'concentration_ability': baseline_player.concentration_difficulty or 0,  # "ability" in baseline
                'motivation_level': baseline_player.motivation_level or 0,
                'cognitive_score': baseline_player.cognitive_test_score or 0,
                'cognitive_reaction_time': baseline_player.cognitive_test_reaction_time or 0
            }
        except:
            baseline_data = {
                'fatigue_level': 0, 'mental_energy': 0, 'concentration_ability': 0,
                'motivation_level': 0, 'cognitive_score': 0, 'cognitive_reaction_time': 0
            }

        # Collect task session data (V1, V2, V3)
        task_sessions_data = []
        task_names = ['Vacancy 1', 'Vacancy 2', 'Vacancy 3']

        for i, round_num in enumerate(task_rounds):
            try:
                round_player = self.player.in_round(round_num)

                def safe_get(field_func, default=0):
                    try:
                        value = field_func()
                        return value if value is not None else default
                    except:
                        return default

                # Get criteria data from HR Coordinator in the same group
                try:
                    hr_coordinator = round_player.group.get_player_by_id(2)
                    criteria_added = safe_get(lambda: hr_coordinator.criteria_added_this_session)
                    criteria_correct = safe_get(lambda: hr_coordinator.criteria_correct_this_session)
                    criteria_incorrect = safe_get(lambda: hr_coordinator.criteria_incorrect_this_session)
                except:
                    criteria_added = criteria_correct = criteria_incorrect = 0

                session_data = {
                    'session': i + 1,
                    'session_name': task_names[i],
                    'role': safe_get(lambda: round_player.selected_role, 'Unknown'),
                    # Individual player data (fatigue, cognitive performance)
                    'fatigue_level': safe_get(lambda: round_player.fatigue_level),
                    'mental_effort': safe_get(lambda: round_player.mental_effort),
                    'concentration_difficulty': safe_get(lambda: round_player.concentration_difficulty),
                    'motivation_level': safe_get(lambda: round_player.motivation_level),
                    'cognitive_score': safe_get(lambda: round_player.cognitive_test_score),
                    'cognitive_reaction_time': safe_get(lambda: round_player.cognitive_test_reaction_time),
                    # Criteria data from HR Coordinator (shared across all players)
                    'criteria_added': criteria_added,
                    'criteria_correct': criteria_correct,
                    'criteria_incorrect': criteria_incorrect,
                }
                task_sessions_data.append(session_data)
            except Exception as e:
                continue

        # Calculate task progression metrics (V1 → V2 → V3)
        if len(task_sessions_data) >= 2:
            # Extract valid values for trend calculations
            fatigue_values = [s['fatigue_level'] for s in task_sessions_data if s['fatigue_level'] > 0]
            cognitive_values = [s['cognitive_score'] for s in task_sessions_data if s['cognitive_score'] > 0]
            effort_values = [s['mental_effort'] for s in task_sessions_data if s['mental_effort'] > 0]
            concentration_values = [s['concentration_difficulty'] for s in task_sessions_data if
                                    s['concentration_difficulty'] > 0]
            motivation_values = [s['motivation_level'] for s in task_sessions_data if s['motivation_level'] > 0]

            # Task progression analysis (first task → last task)
            task_fatigue_increase = fatigue_values[-1] - fatigue_values[0] if len(fatigue_values) >= 2 else 0
            task_cognitive_decline = cognitive_values[0] - cognitive_values[-1] if len(cognitive_values) >= 2 else 0
            task_effort_increase = effort_values[-1] - effort_values[0] if len(effort_values) >= 2 else 0
            task_motivation_change = motivation_values[-1] - motivation_values[0] if len(motivation_values) >= 2 else 0

            # Averages across task sessions only
            average_task_fatigue = sum(fatigue_values) / len(fatigue_values) if fatigue_values else 0
            average_task_effort = sum(effort_values) / len(effort_values) if effort_values else 0
            average_task_concentration = sum(concentration_values) / len(
                concentration_values) if concentration_values else 0
            average_task_motivation = sum(motivation_values) / len(motivation_values) if motivation_values else 0
        else:
            task_fatigue_increase = task_cognitive_decline = task_effort_increase = task_motivation_change = 0
            average_task_fatigue = average_task_effort = average_task_concentration = average_task_motivation = 0

        # Check if current player is HR Coordinator (Player 2)
        is_hr_coordinator = self.player.id_in_group == 2

        result = {
            # Baseline data (separate section)
            'baseline_data': baseline_data,

            # Task sessions data for progression analysis
            'task_sessions': task_sessions_data,
            'total_task_sessions': len(task_sessions_data),

            # Task progression metrics (V1 → V2 → V3)
            'task_fatigue_increase': task_fatigue_increase,
            'task_cognitive_decline': task_cognitive_decline,
            'task_effort_increase': task_effort_increase,
            'task_motivation_change': task_motivation_change,

            # Task session averages
            'average_task_fatigue': average_task_fatigue,
            'average_task_effort': average_task_effort,
            'average_task_concentration': average_task_concentration,
            'average_task_motivation': average_task_motivation,

            # Individual task session values for detailed comparison
            'v1_fatigue': fatigue_values[0] if len(fatigue_values) >= 1 else 0,
            'v2_fatigue': fatigue_values[1] if len(fatigue_values) >= 2 else 0,
            'v3_fatigue': fatigue_values[2] if len(fatigue_values) >= 3 else 0,

            'v1_cognitive': cognitive_values[0] if len(cognitive_values) >= 1 else 0,
            'v2_cognitive': cognitive_values[1] if len(cognitive_values) >= 2 else 0,
            'v3_cognitive': cognitive_values[2] if len(cognitive_values) >= 3 else 0,

            # Inter-vacancy changes
            'v1_to_v2_fatigue_change': fatigue_values[1] - fatigue_values[0] if len(fatigue_values) >= 2 else 0,
            'v2_to_v3_fatigue_change': fatigue_values[2] - fatigue_values[1] if len(fatigue_values) >= 3 else 0,

            'v1_to_v2_cognitive_change': cognitive_values[0] - cognitive_values[1] if len(cognitive_values) >= 2 else 0,
            'v2_to_v3_cognitive_change': cognitive_values[1] - cognitive_values[2] if len(cognitive_values) >= 3 else 0,

            # Role-specific flag
            'is_hr_coordinator': is_hr_coordinator,

            'show_next_button': False,
            'is_final_results': True,
        }
        return result


page_sequence = [
    Consent,
    WaitForVacancy,
    Recruiter,
    HRCoordinator,
    BusinessPartner,
    SelfAssessment,
    CognitiveTestInstructions,
    CognitiveTest,
    CognitiveTestResults,
    FinalResults
]
